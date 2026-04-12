#!/usr/bin/env python3
"""
Local OCR server. Extension'dan gelen screenshotları okur, Word'e yazar.
Çalıştır: python3 server.py
"""
import base64
import io
import os
from datetime import datetime
from flask import Flask, request, jsonify
from PIL import Image
import pytesseract
from docx import Document

app = Flask(__name__)
BASE_DIR = os.path.dirname(__file__)

# Aktif çıktı dosyası (reset/load ile değişir)
current_file = os.path.join(BASE_DIR, "output.docx")


@app.after_request
def add_cors(response):
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Allow-Headers"] = "Content-Type"
    response.headers["Access-Control-Allow-Methods"] = "POST, OPTIONS"
    return response


def get_doc():
    if os.path.exists(current_file):
        return Document(current_file)
    return Document()


@app.route("/ocr", methods=["POST", "OPTIONS"])
def ocr():
    if request.method == "OPTIONS":
        return "", 204
    try:
        data = request.json
        page_num = data.get("page", "?")
        image_b64 = data.get("image", "")

        image_data = base64.b64decode(image_b64.split(",")[-1])
        image = Image.open(io.BytesIO(image_data))

        text = pytesseract.image_to_string(image, lang="tur").strip()

        doc = get_doc()
        doc.add_heading(f"Page {page_num}", level=2)
        doc.add_paragraph(text)
        doc.save(current_file)

        print(f"[Page {page_num}] {len(text)} chars → {os.path.basename(current_file)}")
        return jsonify({"ok": True, "chars": len(text)})
    except Exception as e:
        print(f"ERROR: {e}")
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/load", methods=["POST", "OPTIONS"])
def load():
    if request.method == "OPTIONS":
        return "", 204
    data = request.json
    file_data = base64.b64decode(data.get("file", ""))
    with open(current_file, "wb") as f:
        f.write(file_data)
    print(f"Loaded file → {os.path.basename(current_file)}")
    return jsonify({"ok": True})


@app.route("/reset", methods=["POST", "OPTIONS"])
def reset():
    global current_file
    if request.method == "OPTIONS":
        return "", 204
    # Her yeni çalıştırmada timestamp'li yeni dosya
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    current_file = os.path.join(BASE_DIR, f"output_{timestamp}.docx")
    print(f"New file: {os.path.basename(current_file)}")
    return jsonify({"ok": True, "filename": os.path.basename(current_file)})


if __name__ == "__main__":
    print(f"Server running on http://localhost:8765")
    print(f"Output dir: {BASE_DIR}")
    app.run(port=8765)
